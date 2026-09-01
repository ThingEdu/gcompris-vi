#!/usr/bin/env python3
"""Tài liệu hướng dẫn từng hoạt động GCompris tiếng Việt, xếp theo đúng các mục
trên màn hình chính, kèm ý nghĩa theo tư duy Papert.

Usage: make_docx.py <acts.json> <out.docx>
"""
import json
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

sys.path.insert(0, __file__.rsplit("/", 1)[0])
from make_xlsx import AGE, maker_of, role_of, subject_of  # noqa: E402
from make_xlsx_hoatdong import papert_meaning  # noqa: E402

TIM = RGBColor(0x5C, 0x4B, 0x9E)
DO = RGBColor(0xC8, 0x40, 0x2F)
XAM = RGBColor(0x55, 0x66, 0x72)

# Đúng thứ tự và cách chia mục của màn hình chính GCompris (Menu.qml)
MUC = [
    ("Máy tính", "computer", [],
     "Mục đầu tiên trẻ nhỏ nên vào: làm quen chuột, bàn phím, và những lệnh lập trình đầu tiên."),
    ("Khám phá", "discovery",
     [("Tư duy logic", "logic"), ("Mỹ thuật", "arts"), ("Âm nhạc", "music")],
     "Ba nhánh: nghĩ theo quy luật, vẽ, và nghe – chơi nhạc."),
    ("Khoa học", "sciences",
     [("Thí nghiệm", "experiment"), ("Lịch sử", "history"), ("Địa lý", "geography")],
     "Nhóm gần tinh thần Papert nhất: phần lớn là vi thế giới có luật vật lý riêng để trẻ nghịch."),
    ("Giải trí", "fun", [],
     "Trò chơi vận động tinh và phản xạ. Dùng làm giờ nghỉ giữa buổi hoặc phần thưởng."),
    ("Toán", "math",
     [("Số đếm", "numeration"), ("Số học", "arithmetic"), ("Đo lường", "measures")],
     "Mục lớn nhất. Cẩn thận: phần lớn là luyện tập, nên buổi nào cũng phải có việc làm tay đi kèm."),
    ("Xếp hình", "puzzle", [],
     "Bài toán cầm nắm được. Chỗ học nằm ở lúc trẻ nói ra cách mình nghĩ, không phải lúc ghép xong."),
    ("Đọc", "reading",
     [("Chữ cái", "letters"), ("Từ ngữ", "words"), ("Từ vựng", "vocabulary")],
     "Mục cần xem kỹ nhất trước khi dùng: bộ từ theo cấp độ âm tiết tiếng Việt chưa soạn xong."),
    ("Chiến thuật", "strategy", [],
     "Cờ và trò đối kháng. Cho chơi cặp rồi bắt trẻ giải thích nước đi của mình."),
]

ROLE_ICON = {
    "Vi thế giới": "VI THẾ GIỚI",
    "Công cụ sáng tạo": "CÔNG CỤ SÁNG TẠO",
    "Đồ vật để suy nghĩ": "ĐỒ VẬT ĐỂ SUY NGHĨ",
    "Luyện tập có phản hồi": "LUYỆN TẬP",
}


def p(doc, text="", size=10.5, bold=False, italic=False, color=None,
      space_before=0, space_after=4, style=None):
    par = doc.add_paragraph(style=style)
    run = par.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    if color is not None:
        run.font.color.rgb = color
    par.paragraph_format.space_before = Pt(space_before)
    par.paragraph_format.space_after = Pt(space_after)
    return par


def field(doc, label, text, color=None):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(2)
    par.paragraph_format.left_indent = Pt(14)
    r1 = par.add_run(label + " ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.name = "Calibri"
    r1.font.color.rgb = XAM if color is None else color
    r2 = par.add_run(text)
    r2.font.size = Pt(10)
    r2.font.name = "Calibri"
    if color is not None:
        r2.font.color.rgb = color
    return par


def build(acts, out):
    acts = [a for a in acts if a["name"] not in ("template", "menu")]
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)

    # ---------------------------------------------------------------- bìa
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("GCOMPRIS TIẾNG VIỆT")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = TIM
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("Hướng dẫn từng hoạt động, xếp theo các mục trên màn hình\nÝ nghĩa giáo dục đọc theo tư duy Papert")
    r.font.size = Pt(13)
    r.font.color.rgb = XAM
    x = doc.add_paragraph()
    x.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = x.add_run("ThingEdu · Làng Maker  —  201 hoạt động  —  github.com/ThingEdu/gcompris-vi")
    r.font.size = Pt(10)
    r.font.color.rgb = XAM
    doc.add_page_break()

    # -------------------------------------------------------- cách đọc
    doc.add_heading("Cách đọc tài liệu này", level=1)
    p(doc, "Mỗi hoạt động có năm dòng. Dòng quan trọng nhất là Ý nghĩa Papert và Nối tay — "
           "hai dòng đó quyết định buổi học có giá trị hay chỉ là thời gian ngồi trước màn hình.")
    for lab, txt in [
        ("Việc trẻ làm", "Trẻ thao tác gì trên màn hình, viết đúng như bản dịch đang chạy."),
        ("Mục tiêu", "Mục tiêu học tập do chính nhóm tác giả GCompris đặt ra, đã dịch."),
        ("Vai trò Papert", "Một trong bốn loại ở dưới."),
        ("Ý nghĩa", "Vì sao hoạt động này đáng cho trẻ chơi, đọc theo tư duy Papert."),
        ("Nối tay", "Việc làm bằng vật thật đi kèm. Với nhóm LUYỆN TẬP thì đây là bắt buộc."),
    ]:
        field(doc, lab + ":", txt)
    p(doc)
    doc.add_heading("Bốn vai trò trong tư duy Papert", level=2)
    p(doc, "Seymour Papert phản đối chính cái mà phần mềm giáo dục hay làm nhất: máy hỏi – trẻ đáp – "
           "máy chấm. Cái ông muốn là vi thế giới: một thế giới nhỏ có luật riêng, trẻ thò tay vào "
           "nghịch, đoán, thử, rồi tự thấy mình sai ở đâu.")
    for name, txt in [
        ("VI THẾ GIỚI", "Trẻ thao tác trên một hệ có luật riêng và tự rút ra quy luật. Cho nghịch tự do "
                        "trước khi giảng; hỏi 'con thử đoán xem chuyện gì xảy ra nếu…' rồi im lặng chờ."),
        ("CÔNG CỤ SÁNG TẠO", "Trẻ làm ra sản phẩm của mình, không có đúng sai. Giao một sản phẩm để làm, "
                             "không giao bài tập; in hoặc chiếu sản phẩm trong buổi Chia sẻ."),
        ("ĐỒ VẬT ĐỂ SUY NGHĨ", "Bài toán cụ thể dùng làm chỗ dựa để nghĩ. Cho chơi cặp rồi bắt trẻ nói ra "
                               "chiến thuật; việc nói ra mới là chỗ học, không phải việc thắng."),
        ("LUYỆN TẬP", "Máy hỏi – trẻ đáp. Papert xếp đây là mức dùng máy tính yếu nhất. Tối đa 10 phút, "
                      "và BẮT BUỘC nối sang việc làm tay cùng nội dung."),
    ]:
        par = doc.add_paragraph()
        par.paragraph_format.space_after = Pt(4)
        r1 = par.add_run(name + " — ")
        r1.bold = True
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = TIM if name != "LUYỆN TẬP" else DO
        r2 = par.add_run(txt)
        r2.font.size = Pt(10.5)
    p(doc)
    p(doc, "Ba hoạt động có cảnh báo CẨN THẬN trong phần ý nghĩa — nội dung gốc không hợp với Việt Nam, "
           "chưa nên dùng để dạy. Xem thêm DOCS/NOI_DUNG_CAN_THIET_KE_LAI.md.",
      bold=True, color=DO)
    p(doc, "Thang sao là của chính GCompris: 1–3 sao vàng cho 2–6 tuổi, 4–6 sao đỏ cho 7 tuổi trở lên.",
      italic=True, color=XAM)
    doc.add_page_break()
    return doc, acts


def viet_hoat_dong(doc, a, seen):
    """Một mục hoạt động. seen giữ những hoạt động đã viết ở mục trước."""
    ten = a["vi_title"] or a["en_title"]
    h = doc.add_heading(level=3)
    r = h.add_run(f"{ten}   {'★' * a['diff']}")
    r.font.size = Pt(12)
    if a["name"] in seen:
        par = doc.add_paragraph()
        par.paragraph_format.left_indent = Pt(14)
        par.paragraph_format.space_after = Pt(6)
        rr = par.add_run(f"(đã mô tả ở mục {seen[a['name']]} — cùng một hoạt động, xuất hiện ở hai mục)")
        rr.italic = True
        rr.font.size = Pt(9.5)
        rr.font.color.rgb = XAM
        return
    role = role_of(a)
    field(doc, "Việc trẻ làm:", a["vi_desc"] or "—")
    field(doc, "Mục tiêu:", a["vi_goal"] or "—")
    field(doc, "Vai trò Papert:", ROLE_ICON[role],
          color=DO if role == "Luyện tập có phản hồi" else TIM)
    ynghia = papert_meaning(a)
    field(doc, "Ý nghĩa:", ynghia, color=DO if ynghia.startswith("CẨN THẬN") else None)
    field(doc, "Nối tay:", maker_of(a))
    field(doc, "Độ tuổi gợi ý:", f"{AGE.get(a['diff'], '—')}  ·  môn {subject_of(a)}")


def build_muc(doc, acts):
    import collections

    by_name = {a["name"]: a for a in acts}
    seen = {}
    tong = collections.Counter()

    def loc(tag):
        return [a for a in acts if tag in a["section"].split()]

    for ten_muc, tag, con, mo_ta in MUC:
        trong_muc = loc(tag)
        if not trong_muc:
            continue
        doc.add_heading(f"{ten_muc}  ({len(trong_muc)} hoạt động)", level=1)
        p(doc, mo_ta, italic=True, color=XAM, space_after=8)
        vai = collections.Counter(role_of(a) for a in trong_muc)
        p(doc, "Phân bố vai trò: " + " · ".join(
            f"{ROLE_ICON[k].lower()} {v}" for k, v in vai.most_common()),
          size=9.5, color=XAM, space_after=10)

        da_viet = set()
        if con:
            for ten_con, tag_con in con:
                nhom = [a for a in trong_muc if tag_con in a["section"].split()]
                if not nhom:
                    continue
                doc.add_heading(f"{ten_muc} › {ten_con}  ({len(nhom)})", level=2)
                for a in sorted(nhom, key=lambda x: (x["diff"], x["vi_title"] or x["en_title"])):
                    viet_hoat_dong(doc, a, seen)
                    da_viet.add(a["name"])
                    seen.setdefault(a["name"], f"{ten_muc} › {ten_con}")
                    tong[a["name"]] += 1
            khac = [a for a in trong_muc if a["name"] not in da_viet]
            if khac:
                doc.add_heading(f"{ten_muc} › Khác  ({len(khac)})", level=2)
                for a in sorted(khac, key=lambda x: (x["diff"], x["vi_title"] or x["en_title"])):
                    viet_hoat_dong(doc, a, seen)
                    seen.setdefault(a["name"], ten_muc)
                    tong[a["name"]] += 1
        else:
            for a in sorted(trong_muc, key=lambda x: (x["diff"], x["vi_title"] or x["en_title"])):
                viet_hoat_dong(doc, a, seen)
                seen.setdefault(a["name"], ten_muc)
                tong[a["name"]] += 1
        doc.add_page_break()

    thieu = [a for a in acts if a["name"] not in tong]
    if thieu:
        doc.add_heading(f"Không nằm trong mục nào trên màn hình  ({len(thieu)})", level=1)
        p(doc, "Những hoạt động này chỉ mở được qua ô tìm kiếm hoặc qua dòng lệnh.",
          italic=True, color=XAM)
        for a in sorted(thieu, key=lambda x: (x["diff"], x["vi_title"] or x["en_title"])):
            viet_hoat_dong(doc, a, seen)
    return len(seen), len(thieu)


if __name__ == "__main__":
    acts = json.load(open(sys.argv[1], encoding="utf-8"))
    doc, acts = build(acts, sys.argv[2])
    n, thieu = build_muc(doc, acts)
    doc.save(sys.argv[2])
    print(f"đã ghi {sys.argv[2]} — {n} hoạt động, {thieu} nằm ngoài các mục")
